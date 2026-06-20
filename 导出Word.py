#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""把题库题目导出为高质量 Word(.docx)：公式→可编辑 Word 公式，图片内嵌，表格还原。
   生成两版：空白试卷版 与 含答案解析版。
   依赖：python-docx latex2mathml mathml2omml lxml（首次由 .command 自动安装）。

   被 启动题库.py 的 /api/export-word 调用，也可独立运行：
     python3 导出Word.py --json 题目_xxx.json --title "九年级·圆 专题练习"
     python3 导出Word.py --source "专题24.1 圆【七大题型】"
"""
import os, re, sys, json, argparse
BASE = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(BASE, 'images')

import latex2mathml.converter as _L
from mathml2omml import convert as _m2o
import lxml.etree as etree
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

_NSM = ('xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"')
CONTENT_W = 6.3   # 正文可用宽度(英寸, A4 1英寸边距)

# ────────────────────────── 公式 ──────────────────────────
from functools import lru_cache

@lru_cache(maxsize=20000)
def _omml_str(latex):
    """LaTeX → OMML 字符串（带缓存）；失败返回 None。"""
    try:
        o = _m2o(_L.convert(latex))
        # mathml2omml 生成的 <m:rad> 缺 radPr/deg，Word 会画成空方框 → 补成隐藏次数的平方根
        o = o.replace('<m:rad><m:e>',
                      '<m:rad><m:radPr><m:degHide m:val="1"/></m:radPr><m:deg/><m:e>')
        return o.replace('<m:oMath>', '<m:oMath %s>' % _NSM, 1)
    except Exception:
        return None

def _omml(latex):
    """LaTeX → OMML 元素（每次新建，便于插入文档）；失败返回 None。"""
    o = _omml_str(latex)
    if not o:
        return None
    try:
        return etree.fromstring(o.encode('utf-8'))
    except Exception:
        return None

_SEG = re.compile(r'(\$[^$]*\$)')

# ── 裸算式 → 公式（保守自动识别）──
AUTO_MATH = True
_ALLOWED = r"A-Za-z0-9 \+\-−*/=＝×÷<>＜＞≤≥≠≈⊥∥∠△⊙≌∽^.,，:°%％π√±'’′()（）²³⁰⁴-⁹₀-ₜ"
_RUN = re.compile('[%s]+' % _ALLOWED)
# 触发转换的数学符号（不含单独连字符/区间号，避免误伤 2022-2023 这类）
_TRIG = re.compile(r'[=＝+×÷/<>＜＞≤≥≠≈⊥∥∠△⊙≌∽°%％√±]')
_NUMUNIT = re.compile(r'\d\s*[A-Za-z]|[A-Za-z]\s*\d')   # 数字紧挨字母：量纲/系数（40cm、3x）
_HASLETTER = re.compile(r'[A-Za-z]')                    # 含拉丁字母：变量/线段/多边形（x、AB、OABC）
_ENGWORD = re.compile(r'[a-z]{4,}')                     # 4+连续小写=英文单词 → 不转（保护 AMC 英文题）
_PURE = re.compile(r'\d+(\.\d+)?')                      # 纯数字（含小数）
_YEAR = re.compile(r'1[5-9]\d\d|20\d\d')               # 像年份的 4 位数 → 不转
_NUMLABEL = re.compile(r'[（(]\d+[）)]')                 # (1) 小题号 → 不转
_BAREMAP = {'＝':'=', '−':'-', '×':'\\times ', '÷':'\\div ', '＜':'<', '＞':'>',
            '≤':'\\le ', '≥':'\\ge ', '≠':'\\neq ', '≈':'\\approx ',
            '∠':'\\angle ', '△':'\\triangle ', '⊙':'\\odot ', '⊥':'\\perp ',
            '∥':'\\parallel ', '≌':'\\cong ', '∽':'\\sim ', '°':'^{\\circ}',
            'π':'\\pi ', '±':'\\pm ', '%':'\\%', '％':'\\%', '，':',', '（':'(', '）':')'}
_SQRT = re.compile(r'√\s*(\d+|[A-Za-z]|\([^()]*\)|（[^（）]*）)')

_FRAC = re.compile(r'(\([^()]*\)|[A-Za-z0-9.]+)\s*/\s*(\([^()]*\)|[A-Za-z0-9.]+)')
def _slash_to_frac(s):
    """把 a/b 写成上下分数 \\frac{a}{b}（去掉操作数外层括号）。"""
    def rep(m):
        a = m.group(1).strip(); b = m.group(2).strip()
        if a[:1] == '(' and a[-1:] == ')': a = a[1:-1]
        if b[:1] == '(' and b[-1:] == ')': b = b[1:-1]
        return '\\frac{%s}{%s}' % (a, b)
    return _FRAC.sub(rep, s)

def _bare_to_latex(s):
    s = s.replace('’', "'").replace('′', "'")   # 各种撇号 → LaTeX prime（A'B）
    s = _SQRT.sub(lambda m: '\\sqrt{%s}' % m.group(1).strip('()（）'), s)
    s = s.replace('√', '\\surd ')           # 余下孤立根号的兜底
    s = _slash_to_frac(s)
    for k, v in _BAREMAP.items():
        s = s.replace(k, v)
    return s

def _maybe_wrap(m):
    s = m.group(0)
    lead = s[:len(s) - len(s.lstrip())]; trail = s[len(s.rstrip()):]
    core = s.strip()
    lp = ''; rp = ''                               # 把多余的首部闭括号 / 尾部开括号移出公式
    while core and core[0] in '）),，':  lp += core[0]; core = core[1:]
    while core and core[-1] in '（(,，': rp = core[-1] + rp; core = core[:-1]
    core = core.strip()
    mk = re.match(r'^[（(][A-Za-z][）)]\s*', core)   # 开头的选项标记 (A) 不并入公式
    if mk:
        lp += mk.group(0); core = core[mk.end():].strip()
    if _ENGWORD.search(core):                       # 英文单词/句子 → 不转
        return s
    if _NUMLABEL.fullmatch(core):                   # (1) 小题号 → 不转
        return s
    hasdigit = re.search(r'\d', core)
    # 坐标/带符号数字：(-2,-1)、A(-5,-4)、-3；含逗号或括号即可，带符号时排除 4 位年份区间
    coordlike = bool(hasdigit) and (re.search(r'[,，()（）]', core)
                                    or (re.search(r'[-−+]', core) and not re.search(r'\d{4}', core)))
    pure_ok = bool(_PURE.fullmatch(core)) and not _YEAR.fullmatch(core)   # 纯数字(非年份)
    # 触发：拉丁字母 / 数学符号 / 数字紧挨字母 / 坐标带符号数字 / 纯数字
    if not (_HASLETTER.search(core) or _TRIG.search(core) or _NUMUNIT.search(core) or coordlike or pure_ok):
        return s
    if not re.search(r'[A-Za-z0-9]', core):         # 至少要有字母或数字
        return s
    if core.count('(') != core.count(')') or core.count('（') != core.count('）'):
        return s                                    # 括号不配对 → 不转
    return f'{lead}{lp}${_bare_to_latex(core)}${rp}{trail}'

def _auto_wrap(text):
    if not AUTO_MATH or not text:
        return text
    out = []
    for seg in _SEG.split(text):
        if seg[:1] == '$' and seg[-1:] == '$':
            out.append(seg)
        else:
            out.append(_RUN.sub(_maybe_wrap, seg))
    return ''.join(out)

def _add_inline(para, text):
    """把一行文本(含 $..$ 公式)写入段落。"""
    text = _auto_wrap(text)
    for seg in _SEG.split(text or ''):
        if not seg:
            continue
        if len(seg) >= 2 and seg[0] == '$' and seg[-1] == '$':
            el = _omml(seg[1:-1])
            if el is not None:
                para._p.append(el)
            else:
                para.add_run(seg)            # 转换失败则保留源码
        else:
            para.add_run(seg)

# ────────────────────────── 图片 ──────────────────────────
def _img_path(ref):
    if not ref:
        return None
    name = ref.split('/')[-1]
    p = os.path.join(IMG_DIR, name)
    return p if os.path.isfile(p) else None

# 图片字节解析器（可插拔）：默认读本地 images/；云端可改成 HTTP 抓取。
from io import BytesIO
IMAGE_BYTES = None   # 形如 func(ref)->bytes|None；云服务里覆盖它

def _img_bytes(ref):
    if IMAGE_BYTES is not None:
        try: return IMAGE_BYTES(ref)
        except Exception: return None
    p = _img_path(ref)
    if not p: return None
    try:
        with open(p, 'rb') as f: return f.read()
    except Exception:
        return None

def _img_dims(ref):
    """返回 (宽英寸, 宽高比 w/h)；失败返回 (None, None)。"""
    data = _img_bytes(ref)
    if not data: return None, None
    try:
        from docx.image.image import Image as _DImg
        im = _DImg.from_file(BytesIO(data))
        w_in = im.px_width / (im.horz_dpi or 96)
        return w_in, (im.px_width / im.px_height if im.px_height else 1)
    except Exception:
        return None, None

def _is_small_figure(ref):
    """按宽高比判断：接近方形/竖向(0.4~1.8)→ 适合浮动缩小+环绕；真正宽幅横图(>1.8)→ 行内。"""
    _, ar = _img_dims(ref)
    return ar is not None and 0.4 <= ar <= 1.8

def _put_picture(run, ref, max_w):
    """插入图片并按真实尺寸等比缩放、绝不放大；返回 InlineShape 或 None。"""
    data = _img_bytes(ref)
    if not data: return None
    pic = run.add_picture(BytesIO(data))
    cap = Inches(max_w)
    if pic.width > cap:
        pic.height = Emu(int(pic.height * cap / pic.width)); pic.width = Emu(int(cap))
    return pic

def _floatify(inline, side='right'):
    """把行内图改成浮动锚定图 + 四周型环绕（可在 Word 里自由拖动、文字绕排）。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    drawing = inline.getparent()
    extent = inline.find(qn('wp:extent')); docPr = inline.find(qn('wp:docPr'))
    frame = inline.find(qn('wp:cNvGraphicFramePr')); graphic = inline.find(qn('a:graphic'))
    a = OxmlElement('wp:anchor')
    for k, v in {'distT':'0','distB':'45720','distL':'114300','distR':'114300','simplePos':'0',
                 'relativeHeight':'251658240','behindDoc':'0','locked':'0','layoutInCell':'1','allowOverlap':'1'}.items():
        a.set(k, v)
    sp = OxmlElement('wp:simplePos'); sp.set('x','0'); sp.set('y','0'); a.append(sp)
    ph = OxmlElement('wp:positionH'); ph.set('relativeFrom','column')
    al = OxmlElement('wp:align'); al.text = side; ph.append(al); a.append(ph)
    pv = OxmlElement('wp:positionV'); pv.set('relativeFrom','paragraph')
    of = OxmlElement('wp:posOffset'); of.text = '0'; pv.append(of); a.append(pv)
    a.append(extent)
    ee = OxmlElement('wp:effectExtent')
    for k in ('l','t','r','b'): ee.set(k, '0')
    a.append(ee)
    wr = OxmlElement('wp:wrapSquare'); wr.set('wrapText','bothSides'); a.append(wr)
    a.append(docPr)
    if frame is not None: a.append(frame)
    a.append(graphic)
    drawing.replace(inline, a)

def _add_image(doc, ref, max_w=2.4):
    """行内居中图（解析/多图场景）。"""
    try:
        para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if _put_picture(para.add_run(), ref, max_w) is None:
            para._p.getparent().remove(para._p)   # 没图就别留空段
    except Exception:
        pass

def _add_image_float(para, ref, max_w=2.0, side='right'):
    """把图浮动锚定到给定段落（题干配图）：四周环绕、可拖动。"""
    try:
        pic = _put_picture(para.add_run(), ref, max_w)
        if pic is not None:
            _floatify(pic._inline, side)
    except Exception:
        pass

# ────────────────────────── 表格(@@TBL) ──────────────────────────
def _add_table(doc, body):
    rows = [ln.split('\t') for ln in body.split('\n')]
    maxc = max(len(r) for r in rows) or 1
    t = doc.add_table(rows=len(rows), cols=maxc); t.style = 'Table Grid'
    for ri, cells in enumerate(rows):
        # 少于满列数时让首格横跨补齐(还原合并)
        span = maxc - len(cells) + 1 if len(cells) < maxc else 1
        ci = 0
        for k, c in enumerate(cells):
            cell = t.cell(ri, ci)
            if k == 0 and span > 1:
                cell = cell.merge(t.cell(ri, ci + span - 1)); ci += span
            else:
                ci += 1
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            c = c.strip()
            if c.startswith('@@IMG:'):
                try: _put_picture(cell.paragraphs[0].add_run(), c[6:], 1.7)
                except Exception: pass
            else:
                _add_inline(cell.paragraphs[0], c)

_TBL = re.compile(r'@@TBL\n([\s\S]*?)\n@@/TBL')

def _add_rich(doc, text, first_para=None):
    """渲染富文本：@@TBL 表格 + 多行 + 行内公式。first_para 用于把首行接到既有段落(如题号后)。"""
    s = str(text or ''); last = 0
    for m in _TBL.finditer(s):
        _add_plain(doc, s[last:m.start()], first_para); first_para = None
        _add_table(doc, m.group(1)); last = m.end()
    _add_plain(doc, s[last:], first_para)

def _add_plain(doc, chunk, first_para=None):
    lines = str(chunk or '').split('\n')
    for i, ln in enumerate(lines):
        if i == 0 and first_para is not None:
            _add_inline(first_para, ln)
        else:
            if ln.strip() == '' and i == len(lines) - 1:
                continue
            _add_inline(doc.add_paragraph(), ln)

# ────────────────────────── 选择题选项拆分 ──────────────────────────
_MARK = re.compile(r'[（(]?\s*([A-E])\s*[.．）)、]')

def _split_choices(text):
    pos = [(m.group(1), m.start(), len(m.group(0))) for m in _MARK.finditer(text)]
    if len(pos) < 2:
        return None
    fa = next((i for i, p in enumerate(pos) if p[0] == 'A'), -1)
    if fa < 0:
        return None
    order = 'ABCDE'; seq = []; exp = 0
    for i in range(fa, len(pos)):
        if exp < len(order) and pos[i][0] == order[exp]:
            seq.append(pos[i]); exp += 1
    if len(seq) < 2:
        return None
    stem = text[:seq[0][1]].strip()
    opts = []
    for i, (lt, at, sk) in enumerate(seq):
        end = seq[i + 1][1] if i + 1 < len(seq) else len(text)
        opts.append((lt, text[at + sk:end].strip()))
    return stem, opts

def _disp_w(s):
    import unicodedata
    return sum(2 if unicodedata.east_asian_width(c) in 'WF' else 1 for c in s)

def _add_choices(doc, opts):
    """ABCD 选项用无边框等宽表格排列：很短→一行等分；放不下→两列(A、B / C、D)；超长→一列。
       按显示宽度(中文算2)+「（A）」前缀判断，避免一行挤不下被迫格内换行。"""
    import math
    n = len(opts)
    eff = max((_disp_w(oc) for _, oc in opts), default=0) + 4   # +4 ≈「（A）」前缀宽度
    cols = n if eff <= 8 else (2 if eff <= 22 else 1)
    cols = max(1, min(cols, n))
    rows = math.ceil(n / cols)
    t = doc.add_table(rows=rows, cols=cols)
    t.autofit = False; t.allow_autofit = False
    w = Inches(CONTENT_W / cols)
    for i, (lt, oc) in enumerate(opts):
        cell = t.cell(i // cols, i % cols)
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
        p.add_run(f'（{lt}）')
        _add_inline(p, oc)
        cell.width = w
    for col in t.columns:
        col.width = w

_QNUM = re.compile(r'^[\s　]*(?:第?\s*\d+\s*[.．、）)题]\s*|[（(]\s*\d+\s*[）)]\s*)')
def _strip_qnum(t):
    return _QNUM.sub('', t or '')

# ────────────────────────── 文档组装 ──────────────────────────
ANSWER_SPACE = {'解答题': 4, '解答压轴题': 5, '证明题': 4, '综合题': 5, '计算题': 3}
# 题型在卷内的细排序：选择→选择压轴→填空→填空压轴→解答→解答压轴→证明→综合…
TYPE_ORDER = ['选择题', '选择压轴题', '填空题', '填空压轴题', '解答题', '解答压轴题',
              '证明题', '综合题', '计算题', '判断题', '其他']
# 每个题型归到哪个「大标题」下（压轴题并入对应大类）
BIG = {'选择题': '选择题', '选择压轴题': '选择题',
       '填空题': '填空题', '填空压轴题': '填空题',
       '解答题': '解答题', '解答压轴题': '解答题', '证明题': '解答题', '综合题': '解答题', '计算题': '解答题',
       '判断题': '判断题'}
BIG_ORDER = ['选择题', '填空题', '解答题', '判断题', '其他']
CHN = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十', '十一', '十二']

def _big(t):
    return BIG.get(t or '其他', '其他')

def _order_key(q, i):
    t = q.get('type') or '其他'; b = _big(t)
    return (BIG_ORDER.index(b) if b in BIG_ORDER else len(BIG_ORDER),
            TYPE_ORDER.index(t) if t in TYPE_ORDER else len(TYPE_ORDER), i)

def order_questions(questions):
    """按 大标题(选择/填空/解答) → 细题型(普通先、压轴后) → 原顺序 排序。"""
    return [q for _, q in sorted(enumerate(questions), key=lambda p: _order_key(p[1], p[0]))]

def _set_cn_font(doc, name='宋体', size=11):
    st = doc.styles['Normal']; st.font.name = name; st.font.size = Pt(size)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), name)

def _render_question(doc, q, idx, with_answers):
    text = _strip_qnum(q.get('text', ''))
    p = doc.add_paragraph()
    p.add_run(f'{idx}．').bold = True
    sc = _split_choices(text) if q.get('type') in ('选择题', '选择压轴题') else None
    if sc:
        stem, opts = sc
        _add_rich(doc, stem, first_para=p)
        _add_choices(doc, opts)
    else:
        _add_rich(doc, text, first_para=p)
    imgs = q.get('images') or []
    if len(imgs) == 1 and _is_small_figure(imgs[0]):
        _add_image_float(p, imgs[0], max_w=1.9)   # 小图 → 浮动靠右、文字环绕、可自由拖动
    else:
        for ref in imgs:                          # 宽幅/大图/多图 → 行内居中，限宽不放大
            _add_image(doc, ref, max_w=4.0)

    if with_answers:
        ans = (q.get('solution') or '').strip()
        ana = (q.get('analysis') or '').strip()
        if ans:
            ap = doc.add_paragraph(); r = ap.add_run('【答案/解答】'); r.bold = True; r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
            _add_rich(doc, ans, first_para=ap)
            for ref in (q.get('sol_images') or []): _add_image(doc, ref)
        if ana and ana != ans:
            ap = doc.add_paragraph(); r = ap.add_run('【解析】'); r.bold = True; r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
            _add_rich(doc, ana, first_para=ap)
            for ref in (q.get('ana_images') or []): _add_image(doc, ref)
    else:
        n = ANSWER_SPACE.get(q.get('type', ''), 0)
        for _ in range(n):
            doc.add_paragraph()
    doc.add_paragraph()   # 题间空行

def build_docx(questions, title, with_answers, out_path):
    doc = Document()
    _set_cn_font(doc)
    h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(title + ('（含答案解析）' if with_answers else '')); r.bold = True; r.font.size = Pt(16)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f'共 {len(questions)} 题'); sr.font.size = Pt(9); sr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()
    # 大标题(选择/填空/解答)分节；压轴题并入对应大类；题号全卷连续
    seq = order_questions(questions)
    from collections import Counter
    bigcnt = Counter(_big(q.get('type')) for q in seq)
    ordn = 0; gi = 0; cur_big = None
    for q in seq:
        b = _big(q.get('type'))
        if b != cur_big:
            cur_big = b
            sec = doc.add_paragraph()
            sr = sec.add_run(f'{CHN[gi] if gi < len(CHN) else gi+1}、{b}（共 {bigcnt[b]} 题）')
            sr.bold = True; sr.font.size = Pt(13); gi += 1
        ordn += 1
        _render_question(doc, q, ordn, with_answers)
    doc.save(out_path)
    return out_path

def export_both(questions, title, out_dir=None):
    """生成空白版 + 含答案版，返回两文件路径。"""
    out_dir = out_dir or BASE
    safe = re.sub(r'[\\/:*?"<>|]', '_', title) or '试卷'
    import time; ts = time.strftime('%Y%m%d_%H%M')
    blank = os.path.join(out_dir, f'{safe}_空白卷_{ts}.docx')
    full  = os.path.join(out_dir, f'{safe}_含答案_{ts}.docx')
    build_docx(questions, title, False, blank)
    build_docx(questions, title, True, full)
    return blank, full

# ────────────────────────── CLI ──────────────────────────
def _load_questions(args):
    if args.json:
        d = json.load(open(args.json, encoding='utf-8'))
        return d.get('questions', d) if isinstance(d, dict) else d
    d = json.load(open(os.path.join(BASE, '题库数据.json'), encoding='utf-8'))
    qs = d['questions']
    if args.source: qs = [q for q in qs if (q.get('source') or '') == args.source]
    if args.grade:  qs = [q for q in qs if (q.get('grade') or '') == args.grade]
    if args.topic:  qs = [q for q in qs if (q.get('topic') or '') == args.topic]
    return qs

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--json'); ap.add_argument('--source'); ap.add_argument('--grade')
    ap.add_argument('--topic'); ap.add_argument('--title', default='练习试卷')
    a = ap.parse_args()
    qs = _load_questions(a)
    if not qs:
        print('没有匹配到题目'); return
    b, f = export_both(qs, a.title)
    print('✅ 已生成：'); print('  ', b); print('  ', f)

if __name__ == '__main__':
    main()
